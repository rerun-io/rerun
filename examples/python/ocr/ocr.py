#!/usr/bin/env python3
"""Example template."""

from __future__ import annotations

import argparse

import rerun as rr  # pip install rerun-sdk
import rerun.blueprint as rrb
import docx
from paddleocr import PaddleOCR, PPStructure
from paddleocr.ppstructure.recovery.recovery_to_doc import sorted_layout_boxes, convert_info_docx
import cv2 as cv2
import pandas as pd
import logging
import os
from enum import Enum
from PIL import Image, ImageDraw, ImageFont
# from paddleocr.ppstructure.predict_system import save_structure_res
# from paddleocr.tools.infer.utility import draw_ocr


import numpy as np
from pathlib import Path


# Supportive Classes

class Color:
    Red = (255, 0, 0)
    Green = (0, 255, 0)
    Blue = (0, 0, 255)
    Yellow = (255, 255, 0)
    Cyan = (0, 255, 255)
    Magenta = (255, 0, 255)
    Purple = (128, 0, 128)


class LayoutType(Enum):
    UNKNOWN = (0, "Unknown", Color.Purple)
    TITLE = (1, "title", Color.Red)
    TEXT = (2, "text", Color.Green)
    FIGURE = (3, "figure", Color.Blue)
    FIGURE_CAPTION = (4, "figure_caption", Color.Yellow)
    TABLE = (5, "table", Color.Cyan)
    TABLE_CAPTION = (6, "table_caption", Color.Magenta)
    REFERENCE = (7, "reference", Color.Purple)

    def __str__(self):
        return str(self.value[1])  # Returns the string part

    @property
    def number(self):
        return self.value[0]  # Returns the numerical identifier

    @property
    def type(self):
        return self.value[1]

    @property
    def color(self):
        return self.value[2]  # Returns the color

    @staticmethod
    def get_class_id(text):
        try:
            return LayoutType[text.upper()].number
        except KeyError:
            logging.warning(f"Invalid layout type")
            return 0

    @staticmethod
    def get_type(text):
        try:
            return LayoutType[text.upper()]
        except KeyError:
            logging.warning(f"Invalid layout type")
            return LayoutType.UNKNOWN

    @classmethod
    def get_annotation(cls):
        return [(layout.value[0], layout.value[1], layout.value[2]) for layout in cls]

    # layout.save(type, new_box)
    # layout.save(type, rrb.VisualBounds2D(x_range=[x_min, x_max], y_range=[y_min, y_max]))

    # rr.log(
    #     f"Image/Layout/{layout_id}",
    #     rr.Boxes2D(
    #         array=new_box,
    #         array_format=rr.Box2DFormat.XYXY,
    #         labels=[type],
    #         class_ids=[LayoutType.get_class_id(type)]
    #     ),
    #     timeless=True
    # )


# Main Class for Layout


class Layout:
    def __init__(self):
        self.counts = {layout_type: 0 for layout_type in LayoutType}
        self.records = {layout_type: [] for layout_type in LayoutType}
        self.recovery = """"""

    def add(self, layout_type, bounding_box, detections=[], table=None, figure=None):
        if layout_type in LayoutType:
            self.counts[layout_type] += 1
            name = f"{layout_type}{self.counts[layout_type]}"
            self.records[layout_type].append({
                'type': layout_type,
                'name': name,
                'bounding_box': bounding_box,
                'detections': detections,
                'table': table
            })

            if layout_type != LayoutType.UNKNOWN:
                # self.recovery += f'\n\n({name.title()})\n\n'
                self.recovery += f'\n\n## [{name.title()}](recording://Image/{layout_type.type.title()}/{name.title()})\n\n'

                if layout_type == LayoutType.TABLE:
                    self.recovery += table
                elif layout_type == LayoutType.FIGURE:
                    if not os.path.exists('Figures'):
                        os.makedirs('Figures')
                    img = Image.fromarray(figure)  # Load the image
                    image_path = f'Figures/{name.title()}.jpg'
                    img.save(image_path)  # Save the image
                    absolute_path = os.path.abspath(image_path)
                    self.recovery += '\n' + f"![{name.title()}]({absolute_path})"
                else:
                    for index, detection in enumerate(detections):
                        self.recovery += f' [{detection["text"]}](recording://Image/{layout_type.type.title()}/{name.title()}/Detections/{index})'
                    # if name.title() == "Text2":
                    #
                    # else:
                    #     for detection in detections:
                    #         self.recovery += '\n' + detection['text']
        else:
            logging.warning(f"Invalid layout type: {layout_type}")

    def get_count(self, layout_type):
        if layout_type in LayoutType:
            return self.counts[layout_type]
        else:
            logging.warning(f"Invalid layout type: {layout_type}")

    def get_records(self):
        return self.records

    def get_layout_data(self, line):
        type = line.get('type')
        box = line.get('bbox')
        x_min = box[0]
        y_min = box[1]
        x_max = box[2]
        y_max = box[3]
        new_box = [x_min, y_min, x_max, y_max]

        layout_type = LayoutType.get_type(type)

        if layout_type == LayoutType.TABLE:
            table = self.get_table_markdown(line)
            self.add(layout_type, new_box, table=table)
        elif layout_type == LayoutType.FIGURE:
            detections = self.get_detections(line)
            img = line.get('img')
            self.add(layout_type, new_box, detections, figure=img)
        else:
            detections = self.get_detections(line)
            self.add(layout_type, new_box, detections)

    @staticmethod
    def get_detections(line):
        detections = []
        results = line.get('res')
        # print(results)
        for i, result in enumerate(results):
            text = result.get('text')
            confidence = result.get('confidence')
            box = result.get('text_region')
            x_min, y_min = box[0]
            x_max, y_max = box[2]
            new_box = [x_min, y_min, x_max, y_max]
            detections.append({
                'id': i,
                'text': text,
                'confidence': confidence,
                'box': new_box
            })

        return detections

    @staticmethod
    def get_table_markdown(line):
        html_table = line.get("res").get("html")
        html_data = pd.read_html(html_table)
        df = pd.DataFrame(html_data[0])
        markdown_table = df.to_markdown()
        # rr.log(f"Table", rr.TextDocument(markdown_table, media_type=rr.MediaType.MARKDOWN), timeless=True)
        return markdown_table


def detect_and_log_layout(img_path):
    layout = Layout()

    ocr_model_pp = PPStructure(show_log=False, recovery=True)
    img = cv2.imread(img_path)
    coloured_image = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)

    height, width, _ = coloured_image.shape

    rr.log("Image", rr.Image(coloured_image))
    result_pp = ocr_model_pp(coloured_image)

    h, w, _ = img.shape
    result_pp = sorted_layout_boxes(result_pp, w)

    for line in result_pp:
        layout.get_layout_data(line)

    rr.log("Recovery",
           rr.TextDocument(layout.recovery, media_type=rr.MediaType.MARKDOWN),
           timeless=True
    )

    rr.log(
        "Image",
        rr.AnnotationContext(LayoutType.get_annotation()),
        timeless=True
    )

    layout_detection_boxes = []
    layout_detection_labels = []
    layout_detection_classes = []

    paths = []
    detections_paths = []
    zoom_paths = []
    zoom_paths_figures = []
    zoom_paths_tables = []
    zoom_paths_texts = []

    for layout_type in LayoutType:
        for index, record in enumerate(layout.records[layout_type]):
            rr.log(
                f"Image/{layout_type.type.title()}/{record['name'].title()}",
                rr.Boxes2D(
                    array=record['bounding_box'],
                    array_format=rr.Box2DFormat.XYXY,
                    labels=[str(layout_type.type)],
                    class_ids=[str(layout_type.number)]
                ),
                rr.AnyValues(
                    name=record['name'].title()
                ),
                timeless=True
            )
            detections_paths.append(f"-Image/{layout_type.type.title()}/{record['name'].title()}/Detections/**")
            paths.append(f"-Image/{layout_type.type.title()}/{record['name'].title()}/**")
            for detection in record['detections']:
                rr.log(
                    f"Image/{layout_type.type.title()}/{record['name'].title()}/Detections/{detection['id']}",
                    rr.Boxes2D(
                        array=detection['box'],
                        array_format=rr.Box2DFormat.XYXY,
                        class_ids=[str(layout_type.number)]
                    ),
                    rr.AnyValues(
                        DetectionID=detection['id'],
                        Text=detection['text'],
                        Confidence=detection['confidence']
                    ),
                    timeless=True
                )
            if layout_type == LayoutType.TABLE:
                rr.log(f"Extracted{record['name']}", rr.TextDocument(record['table'], media_type=rr.MediaType.MARKDOWN),
                       timeless=True)
            # layout_detection_boxes.append(record['bounding_box'])
            # layout_detection_labels.append(layout_type.type)
            # layout_detection_classes.append(layout_type.number)

    # valid_items = [item for item in zoom_paths if isinstance(item, ExpectedType)]
    # horizontal_view = rrb.Horizontal(
    #     *valid_items
    # )

    for layout_type in LayoutType:
        for index, record in enumerate(layout.records[layout_type]):
            if layout_type == LayoutType.FIGURE or layout_type == LayoutType.TABLE or layout_type == LayoutType.TEXT:
                current_paths = paths.copy()
                current_paths.remove((f"-Image/{layout_type.type.title()}/{record['name'].title()}/**"))
                box = record['bounding_box']
                x_min, y_min, x_max, y_max = box[0], box[1], box[2], box[3]
                bounds = rrb.VisualBounds2D(x_range=[x_min - 10, x_max + 10], y_range=[y_min - 10, y_max + 10])
                zoom_paths.append(rrb.Spatial2DView(
                    name=record['name'].title(),
                    contents=["Image/**"] + current_paths,
                    visual_bounds=bounds)
                )
                if layout_type == LayoutType.FIGURE:
                    zoom_paths_figures.append(rrb.Spatial2DView(
                        name=record['name'].title(),
                        contents=["Image/**"] + current_paths,
                        visual_bounds=bounds
                    ))
                elif layout_type == LayoutType.TABLE:
                    zoom_paths_tables.append(rrb.Spatial2DView(
                        name=record['name'].title(),
                        contents=["Image/**"] + current_paths,
                        visual_bounds=bounds
                    ))
                elif layout_type == LayoutType.TEXT or layout_type == LayoutType.FIGURE_CAPTION or layout_type == LayoutType.REFERENCE:
                    zoom_paths_texts.append(rrb.Spatial2DView(
                        name=record['name'].title(),
                        contents=["Image/**"] + current_paths,
                        visual_bounds=bounds
                    ))

    tabs = []
    if len(zoom_paths_figures) > 0:
        tabs.append(rrb.Tabs(name="Figures", *zoom_paths_figures))
    if len(zoom_paths_tables) > 0:
        tabs.append(rrb.Tabs(name="Tables", *zoom_paths_tables))
    if len(zoom_paths_texts) > 0:
        tabs.append(rrb.Tabs(name="Texts", *zoom_paths_texts))

    rr.send_blueprint(
        rrb.Blueprint(
            rrb.Vertical(
                rrb.Horizontal(
                    rrb.Spatial2DView(name="Layout", origin='Image/', contents=["Image/**"] + detections_paths),
                    rrb.Spatial2DView(name="Detections", contents=["Image/**"]),
                    rrb.TextDocumentView(name="Recovery", contents='Recovery')
                ),
                rrb.Horizontal(
                    *tabs
                ),
                row_shares=[4, 3],
            ),
            collapse_panels=True,
        )
    )

    # rr.log(
    #     "Image",
    #     rr.AnnotationContext(LayoutType.get_annotation()),
    #     timeless=True
    # )

    # print("FINAL DOC" + str(final_doc))
    # final_doc = ""
    # detections = ""

    # rr.log(f"Detections", rr.TextDocument(detections, media_type=rr.MediaType.MARKDOWN), timeless=True)
    # rr.log(f"Recovery", rr.TextDocument(final_doc, media_type=rr.MediaType.MARKDOWN), timeless=True)
    #
    # my_contents = ["Image/**"]
    # my_lists = []
    #
    # for index, line in enumerate(result_pp):
    #     my_contents.append(f"-Image/Layout/{index}/Detections/**")
    # my_lists.append(rrb.Spatial2DView(name="Output", contents=["Image/**"],
    #                                   visual_bounds=bounds_list[index]))

    # rr.send_blueprint(
    #     rrb.Blueprint(
    #         rrb.Vertical(
    #             rrb.Horizontal(
    #                 rrb.Spatial2DView(name="Layout", origin='Image/', contents=my_contents),
    #                 rrb.Spatial2DView(name="Input", contents=["Image/**"]),
    #                 # rrb.TextDocumentView(name="Recovery", contents='Recovery')
    #             ),
    #             rrb.Horizontal(
    #                 rrb.Spatial2DView(name="Figure1", contents=["Image/**"], visual_bounds=bounds_list[0]),
    #                 rrb.Spatial2DView(name="Figure2", contents=["Image/**"], visual_bounds=bounds_list[1]),
    #                 rrb.Tabs(
    #                     rrb.Spatial2DView(name="Table", contents=["Image/**"], visual_bounds=bounds_list[12]),
    #                     rrb.TextDocumentView(name="Table 1", origin="Table"),
    #                 )
    #             ),
    #             row_shares=[4, 2],
    #         ),
    #         collapse_panels=True,
    #     )
    # )


def get_downloaded_path(dataset_dir: Path, video_name: str) -> str:
    video_file_name = f"{video_name}.mp4"
    destination_path = dataset_dir / video_file_name
    if destination_path.exists():
        logging.info("%s already exists. No need to download", destination_path)
        return str(destination_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Example of using the Rerun visualizer")

    parser.add_argument(
        "--image",
        type=str,
        default="test1",
        choices=["test1", "test2", "layout"],
        help="The example image to run on.",
    )

    parser.add_argument(
        "--language",
        type=str,
        default="en",
        choices=["en", "de", "gr"],
        help="The example image to run on.",
    )

    rr.script_add_args(parser)
    args = parser.parse_args()

    rr.script_setup(
        args,
        "rerun_ocr_example",
        default_blueprint=rrb.Blueprint(
            rrb.Vertical(
                rrb.Spatial2DView(name="Input", contents=["Image/**"]),
            ),
            collapse_panels=True
        )
    )
    rr.script_teardown(args)

    rr.send_blueprint(rrb.Blueprint(
        rrb.Vertical(
            rrb.Spatial2DView(name="Input", contents=["Image/**"]),
        ),
        collapse_panels=True
    ))

    logging.getLogger().addHandler(rr.LoggingHandler("logs/handler"))
    logging.getLogger().setLevel(-1)

    img_path = f"dataset/{args.image}.png"
    detect_and_log_layout(img_path)


if __name__ == "__main__":
    main()

# def handle_text(line, layout_id):
#     detections = []
#     results = line.get('res')
#     box = line.get('bbox')
#     for i, result in enumerate(results):
#         text = result.get('text')
#         confidence = result.get('confidence')
#         box = result.get('text_region')
#         x_min, y_min = box[0]
#         x_max, y_max = box[2]
#         # new_box = [x_min-layout_x_min, y_min-layout_y_min, x_max-layout_x_min, y_max-layout_y_min]
#         new_box = [x_min, y_min, x_max, y_max]
#         detections.append({
#             'text': text,
#             'confidence': confidence,
#             'box': box
#         })
#         # rr.log(
#         #     f"Image/Layout/{layout_id}/Detections/{i}",
#         #     rr.Boxes2D(
#         #         array=new_box,
#         #         array_format=rr.Box2DFormat.XYXY,
#         #         # labels=[text]
#         #     ),
#         #     rr.AnyValues(DetectionID=str(i), Text=text, Confidence=confidence),
#         #     timeless=True
#         # )
#         # # Text=text,
#         #
#         # detections.append(f"{i}. {text} , {confidence}")
#         # final_doc_lines.append(text)
#
#     # # Log the complete document
#     # final_doc = "\n".join(final_doc_lines)
#     # rr.log(f"Output{layout_id}", rr.TextDocument(final_doc, media_type=rr.MediaType.TEXT), timeless=True)
#     return detections
# def handle_table(line, layout_id):
#     rr.log(f"Layouts/{layout_id}", rr.Image(line.get('img')))
#     html_table = line.get("res").get("html")
#     html_data = pd.read_html(html_table)
#     df = pd.DataFrame(html_data[0])
#     markdown_table = df.to_markdown()
#     rr.log(f"Table", rr.TextDocument(markdown_table, media_type=rr.MediaType.MARKDOWN), timeless=True)
#     return str(markdown_table)
#     # rr.log(f"Table{table_id}", rr.TextDocument(markdown_table, media_type=rr.MediaType.MARKDOWN), timeless=True)
#
# def detect_text_from_id1(img_path, lang):
#     # Initialize the OCR model
#     ocr_model = PaddleOCR(lang=lang, use_gpu=False)
#
#     # Read and convert the image
#     img = cv2.imread(img_path)
#     coloured_image = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
#
#     # Log the image
#     rr.log("Image", rr.Image(coloured_image), timeless=True)
#
#     # Perform OCR on the image
#     result = ocr_model.ocr(coloured_image)
#
#     boxes = [line[0] for line in result[0]]
#     txts = [line[1][0] for line in result[0]]
#     scores = [line[1][1] for line in result[0]]
#
#     height, width, _ = img.shape
#     white_canvas = np.ones((height, width, 3), dtype=np.uint8) * 255
#
#     # image_pil = Image.fromarray(cv2.cvtColor(img, cv2.COLOR_BGR2RGB))
#     image_pil = Image.fromarray(cv2.cvtColor(white_canvas, cv2.COLOR_BGR2RGB))
#     draw = ImageDraw.Draw(image_pil)
#     font = ImageFont.truetype("latin.ttf", 16)  # You can use any TTF font file available in your system
#
#     for line in result[0]:
#         box = line[0]
#         text = line[1][0]
#         score = line[1][1]
#
#         # Draw the bounding box
#         pts = np.array(box, np.int32)
#         pts = pts.reshape((-1, 1, 2))
#         cv2.polylines(img, [pts], isClosed=True, color=(0, 255, 0), thickness=2)
#
#         # Put the detected text
#         (x, y) = box[0]
#
#         draw.text((x, y), text, font=font, fill=(0, 0, 0, 255))  # White text
#
#     # Convert image back to OpenCV format
#     img_with_text = cv2.cvtColor(np.array(image_pil), cv2.COLOR_RGB2BGR)
#
#     rr.log("MyImage", rr.Image(img_with_text))
#
#     # Initialize containers for extracted information
#     boxes, texts, scores = [], [], []
#     final_doc_lines = []
#
#     # Extract information from OCR results
#     for i, res in enumerate(result[0]):
#         box = res[0]
#         text = res[1][0]
#         score = res[1][1]
#
#         x_min, y_min = box[0]
#         x_max, y_max = box[2]
#         new_box = [x_min, y_min, x_max, y_max]
#
#         # Log detection details
#         rr.log(
#             f"Image/Detections/{i}",
#             rr.Boxes2D(array=new_box, array_format=rr.Box2DFormat.XYXY, labels=[text]),
#             rr.AnyValues(DetectionID=str(i), Text=text, Score=score),
#             timeless=True
#         )
#
#         texts.append(text)
#         scores.append(score)
#         final_doc_lines.append(f"{i}. {text} \t ({score})")
#         logging.info("Detection %d with text: %s, and score: %.2f", i, text, score)
#
#     # Log the complete document
#     final_doc = "\n".join(final_doc_lines)
#     rr.log("Detections", rr.TextDocument(final_doc, media_type=rr.MediaType.TEXT), timeless=True)
#
#     # Extract and clean specific text fields for final document
#     text_birth = texts[12].removesuffix("DEUTSCH")
#
#     final_doc_id = f'''
#     ID Number:  {texts[2]}, {scores[2]}
#     Surname: {texts[5]}, {scores[5]}
#     Name:  {texts[6]}, {scores[6]}
#     Date of Birthday: {text_birth}, {scores[12]}
#     Place of Birth: {texts[14]}, {scores[14]}
#     Date of Expiry: {texts[16]}, {scores[16]}
#     '''
#
#     # Log the final document details
#     rr.log("IDDetails", rr.TextDocument(final_doc_id.strip(), media_type=rr.MediaType.TEXT), timeless=True)
#
#
# def detect_text_from_id(img_path, lang):
#     ocr_model = PaddleOCR(lang=lang, use_gpu=False)
#     img = cv2.imread(img_path)
#     coloured_image = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
#     rr.log("Image", rr.Image(coloured_image), timeless=True)
#     result = ocr_model.ocr(coloured_image)
#
#     final_doc = ""
#
#     # extract information from result
#     boxes, texts, scores = [], [], []
#     for i, res in enumerate(result[0]):
#         boxes.append(res[0])
#
#         box = res[0]
#         x_min = box[0][0]
#         y_min = box[0][1]
#         x_max = box[2][0]
#         y_max = box[2][1]
#         new_box = [x_min, y_min, x_max, y_max]
#
#         text = res[1][0]
#         score = res[1][1]
#
#         rr.log(
#             f"Image/Detections/{i}",
#             rr.Boxes2D(
#                 array=new_box,
#                 array_format=rr.Box2DFormat.XYXY
#             ),
#             rr.AnyValues(**{"DetectionID": str(i), "Text": text, "Score": score}),
#             timeless=True
#         )
#
#         texts.append(text)
#         scores.append(score)
#         final_doc = final_doc + f"\nID: {i},\t Text: {text},\t Score: {score}"
#         logging.info("Detection %d with text: %s, and score: %.2f", i, text, score)
#
#     rr.log("Detections", rr.TextDocument(final_doc.strip(), media_type=rr.MediaType.TEXT), timeless=True)
#
#     text_birth = str(texts[12])
#     text_birth = text_birth.removesuffix("DEUTSCH")
#
#     final_doc_id = f'''
# ID Number:  {texts[2]}, {scores[2]}
# Surname: {texts[5]}, {scores[5]}
# Name:  {texts[6]}, {scores[6]}
# Date of Birthday: {text_birth}, {scores[12]}
# Place of Birth: {texts[14]}, {scores[14]}
# Date of Expiry: {texts[16]}, {scores[16]}
# '''
#
#     rr.log("IDDetails", rr.TextDocument(final_doc_id.strip(), media_type=rr.MediaType.TEXT), timeless=True)
